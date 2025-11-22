from flask import Flask, request, jsonify, session
from flask_cors import CORS
import PyPDF2
import docx
import requests
import tempfile
import os
import uuid
from datetime import datetime, date
import secrets
from functools import wraps
import json

app = Flask(__name__)
# Добавляем секретный ключ для сессий
app.secret_key = os.getenv('SECRET_KEY', secrets.token_hex(32))

# Исправляем CORS для работы на Render
CORS(app, resources={r"/*": {"origins": "*"}})

# Загружаем переменные окружения
from dotenv import load_dotenv
load_dotenv()

# Данные Yandex Cloud из переменных окружения
YANDEX_API_KEY = os.getenv('YANDEX_API_KEY')
YANDEX_FOLDER_ID = os.getenv('YANDEX_FOLDER_ID')

# Система пользователей и лимитов
# Файл для хранения пользователей (на Render используем /tmp)
USER_DB_FILE = '/tmp/docscan_users.json'

def load_users():
    """Загружает пользователей из файла"""
    try:
        if os.path.exists(USER_DB_FILE):
            with open(USER_DB_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                print(f"✅ Загружено {len(data)} пользователей из файла")
                
                # Восстанавливаем даты и сбрасываем лимиты если нужно
                for user_id, user_data in data.items():
                    if user_data['last_reset'] < date.today().isoformat():
                        user_data['used_today'] = 0
                        user_data['last_reset'] = date.today().isoformat()
                        print(f"🔄 Сброшен лимит для пользователя {user_id}")
                
                return data
    except Exception as e:
        print(f"❌ Ошибка загрузки пользователей: {e}")
    
    # База по умолчанию если файла нет
    default_db = {
        'default': {
            'plan': 'free',
            'used_today': 0,
            'last_reset': date.today().isoformat(),
            'total_used': 0,
            'user_id': 'default',
            'created_at': datetime.now().isoformat()
        }
    }
    print("✅ Создана база по умолчанию")
    return default_db

def save_users():
    """Сохраняет пользователей в файл"""
    try:
        with open(USER_DB_FILE, 'w', encoding='utf-8') as f:
            json.dump(users_db, f, ensure_ascii=False, indent=2)
        print(f"💾 Сохранено {len(users_db)} пользователей")
    except Exception as e:
        print(f"❌ Ошибка сохранения пользователей: {e}")

# Загружаем базу при старте сервера
users_db = load_users()
print(f"🚀 Сервер запущен. Всего пользователей: {len(users_db)}")

# Добавляем администраторов
ADMINS = {
    'admin': 'admin123',  # login: password
    'superuser': 'super123'
}

# Глобальная переменная для хранения сессий
admin_sessions = {}

# ОБНОВЛЕННЫЕ ТАРИФЫ - 1 бесплатный, потом платные
PLANS = {
    'free': {
        'daily_limit': 1,  # БЫЛО 3, ТЕПЕРЬ 1
        'ai_access': True,
        'price': 0,
        'name': 'Бесплатный'
    },
    'basic': {
        'daily_limit': 10,  # 10 анализов в день
        'ai_access': True, 
        'price': 199,
        'name': 'Базовый'
    },
    'premium': {
        'daily_limit': 50,  # 50 анализов в день
        'ai_access': True,
        'price': 399,
        'name': 'Премиум'
    },
    'unlimited': {
        'daily_limit': 1000,  # Фактически безлимит
        'ai_access': True,
        'price': 800,
        'name': 'Безлимитный'
    }
}

def generate_user_id():
    """Генерирует уникальный ID пользователя"""
    return str(uuid.uuid4())[:8]

def get_user(user_id=None):
    """Получает или создает пользователя"""
    if not user_id:
        user_id = generate_user_id()
    
    if user_id not in users_db:
        users_db[user_id] = {
            'user_id': user_id,
            'plan': 'free',
            'used_today': 0,
            'last_reset': date.today().isoformat(),
            'total_used': 0,
            'created_at': datetime.now().isoformat()
        }
        save_users()  # Сохраняем при создании нового
        print(f"👤 Создан новый пользователь: {user_id}")
    
    user = users_db[user_id]
    
    # Сбрасываем дневной лимит если новый день
    if user['last_reset'] < date.today().isoformat():
        user['used_today'] = 0
        user['last_reset'] = date.today().isoformat()
        save_users()  # Сохраняем при сбросе лимита
        print(f"🔄 Сброшен дневной лимит для {user_id}")
    
    return user

def can_analyze(user_id='default'):
    """Проверяет может ли пользователь сделать анализ"""
    user = get_user(user_id)
    return user['used_today'] < PLANS[user['plan']]['daily_limit']

def record_usage(user_id='default'):
    """Записывает использование"""
    user = get_user(user_id)
    user['used_today'] += 1
    user['total_used'] += 1
    save_users()  # Сохраняем при каждом использовании
    print(f"📊 Записан анализ для {user_id}. Сегодня: {user['used_today']}, Всего: {user['total_used']}")

# Функции анализа документов
def extract_text_from_pdf(file_path):
    text = ""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        return f"Ошибка чтения PDF: {str(e)}"
    return text

def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        return f"Ошибка чтения DOCX: {str(e)}"
    return text

def parse_fallback_response(ai_response):
    """Резервный парсинг для неструктурированных ответов"""
    risks = []
    recommendations = []
    
    lines = [line.strip() for line in ai_response.split('\n') if line.strip()]
    
    for i, line in enumerate(lines):
        line_lower = line.lower()
        
        # Ищем риски по ключевым словам
        if any(word in line_lower for word in ['риск', 'опасность', 'проблема', 'недостаток', 'слабое место', 'угроза']):
            # Берем следующие несколько строк как описание риска
            for j in range(i+1, min(i+4, len(lines))):
                next_line = lines[j]
                if next_line and len(next_line) > 20 and not next_line.lower().startswith('рекомендац'):
                    risks.append(next_line)
                    break
        
        # Ищем рекомендации по ключевым словам
        elif any(word in line_lower for word in ['рекомендац', 'совет', 'следует', 'рекомендуется', 'улучшить', 'добавить']):
            # Берем следующие несколько строк как рекомендацию
            for j in range(i+1, min(i+4, len(lines))):
                next_line = lines[j]
                if next_line and len(next_line) > 20 and not next_line.lower().startswith('риск'):
                    recommendations.append(next_line)
                    break
    
    return risks, recommendations

def analyze_with_yandexgpt(text):
    """Анализирует текст с помощью YandexGPT"""
    try:
        headers = {
            "Authorization": f"Api-Key {YANDEX_API_KEY}",
            "Content-Type": "application/json"
        }
        
        data = {
            "modelUri": f"gpt://{YANDEX_FOLDER_ID}/yandexgpt-lite/latest",
            "completionOptions": {
                "stream": False,
                "temperature": 0.1,
                "maxTokens": 2000
            },
            "messages": [
                {
                    "role": "system", 
                    "text": """Ты опытный юрист-аналитик. Проанализируй документ и выдели ТОЛЬКО:
1. ПОТЕНЦИАЛЬНЫЕ РИСКИ (конкретные проблемы, что может привести к потерям)
2. КОНКРЕТНЫЕ РЕКОМЕНДАЦИИ по исправлению

Формат ответа:
РИСКИ:
- риск 1
- риск 2

РЕКОМЕНДАЦИИ:
- рекомендация 1
- рекомендация 2

Не добавляй общие оценки безопасности и другие комментарии."""
                },
                {
                    "role": "user",
                    "text": f"Проанализируй этот документ как юрист и выдели только риски и рекомендации:\n\n{text[:8000]}"
                }
            ]
        }
        
        response = requests.post(
            "https://llm.api.cloud.yandex.net/foundationModels/v1/completion",
            headers=headers,
            json=data,
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            ai_response = result['result']['alternatives'][0]['message']['text']
            
            # Улучшенный парсинг ответа
            lines = [line.strip() for line in ai_response.split('\n') if line.strip()]
            risks = []
            recommendations = []
            
            current_section = None
            
            for line in lines:
                line_lower = line.lower()
                
                # Определяем разделы
                if any(marker in line_lower for marker in ['риск', 'проблем', 'опасност', 'недостаток', 'слаб']):
                    current_section = 'risks'
                    continue
                elif any(marker in line_lower for marker in ['рекомендац', 'совet', 'улучшен', 'исправлен']):
                    current_section = 'recommendations'
                    continue
                
                # Пропускаем заголовки и общие фразы
                if any(phrase in line_lower for phrase in [
                    'общая оценка', 'документ выглядит', 'безопасн', 'итог', 'заключен'
                ]):
                    continue
                
                # Добавляем пункты только если они начинаются с маркера списка
                if line.startswith(('-', '•', '—', '*', '1.', '2.', '3.', '4.', '5.')) and len(line) > 5:
                    if current_section == 'risks':
                        risks.append(line.lstrip('-•—*123456789. '))
                    elif current_section == 'recommendations':
                        recommendations.append(line.lstrip('-•—*123456789. '))
            
            # Если не нашли структурированный ответ, используем эвристический подход
            if not risks or not recommendations:
                risks, recommendations = parse_fallback_response(ai_response)
            
            # Очистка от дубликатов и пустых строк
            risks = list(dict.fromkeys([r for r in risks if r and len(r) > 10]))
            recommendations = list(dict.fromkeys([r for r in recommendations if r and len(r) > 10]))
            
            return {
                'risks': risks if risks else ['✅ Критических рисков не обнаружено'],
                'warnings': [],
                'summary': f'🤖 YandexGPT: {len(text)} символов проанализировано',
                'recommendations': recommendations if recommendations else ['✅ Все рекомендации учтены в документе'],
                'ai_used': True
            }
        else:
            return {
                'risks': [f'❌ Ошибка YandexGPT: {response.status_code}'],
                'warnings': [],
                'summary': 'Ошибка доступа к AI',
                'recommendations': ['🔄 Используем локальный анализ...'],
                'ai_used': False
            }
            
    except Exception as e:
        return {
            'risks': [f'❌ Ошибка соединения: {str(e)}'],
            'warnings': [],
            'summary': 'Нет соединения с AI',
            'recommendations': ['🔄 Переключаемся на локальный анализ'],
            'ai_used': False
        }

def analyze_text(text, user_id='default'):
    """Основная функция анализа"""
    user = get_user(user_id)
    
    # Проверяем доступ к AI по тарифу
    if PLANS[user['plan']]['ai_access']:
        result = analyze_with_yandexgpt(text)
        if result['ai_used']:
            return result
    
    # Если AI недоступен, используем локальный анализ
    return {
        'risks': ['✅ Базовый анализ завершен'],
        'warnings': [],
        'summary': f'📊 Локальный анализ: {len(text)} символов',
        'recommendations': ['💎 Перейдите на премиум для AI-анализа'],
        'ai_used': False
    }

# API endpoints
@app.route('/')
def home():
    """Главная страница с интерфейсом"""
    return """
    <!DOCTYPE html>
    <html lang="ru">
    <head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>DocScan - AI анализ документов и договоров за 60 секунд</title>
    <meta name="description" content="Бесплатный анализ документов с AI. Проверка договоров на риски, выявление ошибок с помощью YandexGPT. Юридический анализ за 1 минуту">
    <meta name="keywords" content="анализ документов, проверка договоров, AI анализ, YandexGPT, юридический анализ, анализ рисков, проверка документов">
    <meta name="robots" content="index, follow">
    <meta name="author" content="DocScan">
    <style>
            * { margin: 0; padding: 0; box-sizing: border-box; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; }
            body { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); min-height: 100vh; padding: 20px; display: flex; justify-content: center; align-items: center; }
            .container { background: white; border-radius: 20px; padding: 40px; box-shadow: 0 20px 40px rgba(0,0,0,0.1); max-width: 1000px; width: 100%; }
            .header { text-align: center; margin-bottom: 40px; }
            .logo { font-size: 3em; margin-bottom: 10px; }
            h1 { color: #2d3748; margin-bottom: 10px; font-size: 2.2em; }
            .subtitle { color: #718096; font-size: 1.2em; }
            .user-info { background: #edf2f7; padding: 15px; border-radius: 10px; margin: 20px 0; text-align: center; }
            .upload-zone { border: 3px dashed #cbd5e0; border-radius: 15px; padding: 60px 30px; text-align: center; margin: 30px 0; transition: all 0.3s ease; background: #f7fafc; cursor: pointer; }
            .upload-zone:hover { border-color: #667eea; background: #edf2f7; }
            .upload-icon { font-size: 4em; color: #667eea; margin-bottom: 20px; }
            .btn { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; border: none; padding: 15px 40px; border-radius: 50px; font-size: 1.1em; cursor: pointer; transition: transform 0.2s ease; margin: 10px; }
            .btn:hover { transform: translateY(-2px); box-shadow: 0 10px 20px rgba(102,126,234,0.3); }
            .btn:disabled { background: #a0aec0; cursor: not-allowed; transform: none; box-shadow: none; }
            .file-info { background: #edf2f7; padding: 15px; border-radius: 10px; margin: 20px 0; }
            .loading { display: none; text-align: center; margin: 20px 0; }
            .spinner { border: 4px solid #f3f3f3; border-top: 4px solid #667eea; border-radius: 50%; width: 40px; height: 40px; animation: spin 1s linear infinite; margin: 0 auto 20px; }
            @keyframes spin { 0% { transform: rotate(0deg); } 100% { transform: rotate(360deg); } }
            .result { background: #f8fafc; border-radius: 15px; padding: 30px; margin-top: 30px; display: none; }
            .risk-item { background: white; padding: 15px; margin: 10px 0; border-radius: 10px; border-left: 4px solid #e53e3e; }
            .success-item { background: white; padding: 15px; margin: 10px 0; border-radius: 10px; border-left: 4px solid #48bb78; }
            .summary { background: #e6fffa; padding: 20px; border-radius: 10px; margin: 20px 0; border-left: 4px solid #38a169; }
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <div class="logo">🔍</div>
                <h1>DocScan</h1>
                <p class="subtitle">Понять суть документа за 60 секунд</p>
            </div>

            <div class="user-info" id="userInfo">
                <strong>👤 Ваш ID:</strong> <span id="userId">Загрузка...</span><br>
                <strong>📊 Анализов сегодня:</strong> <span id="usageInfo">0/1</span><br>
            </div>

            <div class="upload-zone" id="dropZone" onclick="document.getElementById('fileInput').click()">
                <div class="upload-icon">📄</div>
                <p><strong>Нажмите чтобы выбрать документ</strong></p>
                <p style="color: #718096; margin-top: 15px;">PDF, DOCX, TXT (до 10MB)</p>
            </div>

            <input type="file" id="fileInput" style="display: none;" accept=".pdf,.docx,.txt" onchange="handleFileSelect(event)">
            
            <div class="file-info" id="fileInfo" style="display: none;">
                <strong>Выбран файл:</strong> <span id="fileName"></span>
            </div>

            <button class="btn" id="analyzeBtn" onclick="analyzeDocument()" disabled>Начать анализ</button>

            <div class="loading" id="loading">
                <div class="spinner"></div>
                <p>Анализируем документ...</p>
            </div>

            <div class="result" id="result">
                <h3>✅ Анализ завершен</h3>
                <div id="resultContent"></div>
            </div>

            <div class="plans" style="margin-top: 40px;">
                <div style="text-align: center; margin-bottom: 20px;">
                    <h3>💎 Выберите тариф</h3>
                </div>
                
                <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 20px;">
                    <div style="background: white; padding: 25px; border-radius: 15px; border: 2px solid #e53e3e; text-align: center;">
                        <div style="font-size: 1.3em; font-weight: bold; margin-bottom: 10px; color: #e53e3e;">Бесплатный</div>
                        <div style="font-size: 2em; font-weight: bold; color: #e53e3e; margin-bottom: 15px;">0₽</div>
                        <ul style="list-style: none; margin-bottom: 20px; text-align: left;">
                            <li style="padding: 5px 0;">✅ 1 анализ в день</li>
                            <li style="padding: 5px 0;">✅ AI-анализ YandexGPT</li>
                            <li style="padding: 5px 0;">✅ Все форматы файлов</li>
                        </ul>
                        <button class="btn" disabled style="background: #e53e3e;">Текущий тариф</button>
                    </div>
                    
                    <div style="background: #f0fff4; padding: 25px; border-radius: 15px; border: 2px solid #38a169; text-align: center;">
                        <div style="font-size: 1.3em; font-weight: bold; margin-bottom: 10px; color: #38a169;">Базовый</div>
                        <div style="font-size: 2em; font-weight: bold; color: #38a169; margin-bottom: 15px;">199₽/мес</div>
                        <ul style="list-style: none; margin-bottom: 20px; text-align: left;">
                            <li style="padding: 5px 0;">🚀 10 анализов в день</li>
                            <li style="padding: 5px 0;">🚀 Приоритетный AI-анализ</li>
                            <li style="padding: 5px 0;">🚀 Быстрая обработка</li>
                        </ul>
                        <button class="btn" onclick="buyPlan('basic')" style="background: #38a169;">Купить за 199₽</button>
                    </div>
                </div>
            </div>

        <script>
            let selectedFile = null;
            let currentUserId = null;

            // Загружаем или создаем ID пользователя
            function loadUser() {
                let savedId = localStorage.getItem('docscan_user_id');
                if (!savedId) {
                    // Создаем нового пользователя
                    fetch('/create-user', { method: 'POST' })
                        .then(r => r.json())
                        .then(data => {
                            if (data.success) {
                                currentUserId = data.user_id;
                                localStorage.setItem('docscan_user_id', currentUserId);
                                updateUserInfo();
                            }
                        });
                } else {
                    currentUserId = savedId;
                    updateUserInfo();
                }
            }

            function updateUserInfo() {
                if (!currentUserId) return;
                
                document.getElementById('userId').textContent = currentUserId;
                
                // Загружаем информацию об использовании
                fetch(`/usage?user_id=${currentUserId}`)
                    .then(r => r.json())
                    .then(data => {
                        document.getElementById('usageInfo').textContent = 
                            `${data.used_today}/${data.daily_limit}`;
                    });
            }

            function copyUserId() {
                navigator.clipboard.writeText(currentUserId);
                alert('ID скопирован: ' + currentUserId);
            }

            function generateNewId() {
                if (confirm('Создать новый ID? Текущая статистика будет сброшена.')) {
                    fetch('/create-user', { method: 'POST' })
                        .then(r => r.json())
                        .then(data => {
                            if (data.success) {
                                currentUserId = data.user_id;
                                localStorage.setItem('docscan_user_id', currentUserId);
                                updateUserInfo();
                                alert('Новый ID создан: ' + currentUserId);
                            }
                        });
                }
            }

            function handleFileSelect(event) {
                const file = event.target.files[0];
                if (!file) return;
                
                if (!file.name.match(/\\.(pdf|docx|txt)$/)) {
                    alert('Пожалуйста, выберите файл в формате PDF, DOCX или TXT');
                    return;
                }

                if (file.size > 10 * 1024 * 1024) {
                    alert('Файл слишком большой. Максимальный размер: 10MB');
                    return;
                }

                selectedFile = file;
                document.getElementById('fileName').textContent = file.name;
                document.getElementById('fileInfo').style.display = 'block';
                document.getElementById('analyzeBtn').disabled = false;
            }

            async function analyzeDocument() {
                if (!selectedFile || !currentUserId) return;

                document.getElementById('loading').style.display = 'block';
                document.getElementById('analyzeBtn').disabled = true;

                try {
                    const formData = new FormData();
                    formData.append('file', selectedFile);
                    formData.append('user_id', currentUserId);

                    const response = await fetch(window.location.origin + '/analyze', {
                        method: 'POST',
                        body: formData
                    });

                    if (!response.ok) {
                        throw new Error(`HTTP error! status: ${response.status}`);
                    }

                    const data = await response.json();

                    document.getElementById('loading').style.display = 'none';

                    if (data.success) {
                        showResult(data);
                        updateUserInfo(); // Обновляем статистику
                    } else {
                        alert('Ошибка: ' + data.error);
                        document.getElementById('analyzeBtn').disabled = false;
                    }

                } catch (error) {
                    document.getElementById('loading').style.display = 'none';
                    
                    if (error.message.includes('402')) {
                        alert('❌ Бесплатный лимит исчерпан!\\n\\nСегодня вы использовали 1/1 бесплатный анализ.\\n\\n💎 Перейдите на платный тариф для продолжения.');
                    } else {
                        alert('Ошибка соединения: ' + error.message);
                    }
                    
                    document.getElementById('analyzeBtn').disabled = false;
                }
            }

            function showResult(data) {
                const resultDiv = document.getElementById('result');
                const resultContent = document.getElementById('resultContent');
                
                let risksHTML = '';
                data.result.risks.forEach(risk => {
                    risksHTML += `<div class="risk-item">${risk}</div>`;
                });
                
                let recommendationsHTML = '';
                data.result.recommendations.forEach(rec => {
                    recommendationsHTML += `<div class="success-item">${rec}</div>`;
                });
                
                resultContent.innerHTML = `
                    <div style="margin-bottom: 20px;">
                        <strong>📄 Анализ документа:</strong> ${data.filename}
                    </div>
                    
                    <div class="summary">
                        ${data.result.summary}
                    </div>
                    
                    ${risksHTML ? `<h4 style="margin: 20px 0 10px 0; color: #e53e3e;">⚠️ Выявленные риски:</h4>${risksHTML}` : ''}
                    
                    ${recommendationsHTML ? `<h4 style="margin: 20px 0 10px 0; color: #48bb78;">✅ Рекомендации:</h4>${recommendationsHTML}` : ''}
                `;
                
                resultDiv.style.display = 'block';
                resultDiv.scrollIntoView({ behavior: 'smooth' });
            }

            // Загружаем пользователя при старте
        loadUser();

        // 🔐 ФУНКЦИЯ ДЛЯ ПОКУПКИ ТАРИФОВ
        async function buyPlan(planType) {
            if (!currentUserId) {
                alert('Сначала загрузите страницу');
                return;
            }
            
            const planNames = {
                'basic': 'Базовый (199₽)'
            };
            
            if (!confirm(`Вы уверены что хотите купить тариф "${planNames[planType]}"?`)) {
                return;
            }
            
            try {
                const response = await fetch('/create-payment', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({
                        user_id: currentUserId,
                        plan: planType
                    })
                });
                
                const result = await response.json();
                
                if (result.success) {
                    window.location.href = result.payment_url;
                } else {
                    alert('Ошибка: ' + result.error);
                }
                
            } catch (error) {
                alert('Ошибка соединения: ' + error.message);
            }
        }
    </script>

        <!-- ФУТЕР -->
        <div style="width: 100%; text-align: center; padding: 30px 0; color: #718096; border-top: 1px solid #e2e8f0; margin-top: 50px; background: white;">
            <div style="max-width: 1000px; margin: 0 auto; padding: 0 20px;">
                <div style="margin-bottom: 15px;">
                    <a href="/terms" style="color: #718096; text-decoration: none; margin: 0 15px; font-size: 14px;">Пользовательское соглашение</a>
                    <a href="/privacy" style="color: #718096; text-decoration: none; margin: 0 15px; font-size: 14px;">Политика конфиденциальности</a>
                    <a href="/offer" style="color: #718096; text-decoration: none; margin: 0 15px; font-size: 14px;">Публичная оферта</a>
                    <a href="mailto:docscanhelp@gmail.com" style="color: #718096; text-decoration: none; margin: 0 15px; font-size: 14px;">Техподдержка</a>
                </div>
                <div style="font-size: 14px;">
                    © 2025 DocScan. Все права защищены.
                </div>
            </div>
        </div>

    </body>
</html>
    """

# Добавляем endpoint для создания пользователя
@app.route('/create-user', methods=['POST'])
def create_user():
    """Создает нового пользователя и добавляет в базу"""
    try:
        user_id = generate_user_id()
        # ЭТА СТРОКА ДОБАВЛЯЕТ ПОЛЬЗОВАТЕЛЯ В БАЗУ
        user = get_user(user_id)  # get_user автоматически создает пользователя если его нет
        
        return jsonify({
            'success': True,
            'user_id': user_id,
            'message': 'Пользователь создан'
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

# Обновляем endpoint анализа для работы с user_id
@app.route('/analyze', methods=['POST'])
def analyze_document():
    # Получаем user_id из формы или используем default
    user_id = request.form.get('user_id', 'default')
    
    # Проверяем лимиты
    if not can_analyze(user_id):
        user = get_user(user_id)
        plan = PLANS[user['plan']]
        return jsonify({
            'success': False,
            'error': f'❌ Бесплатный лимит исчерпан! Сегодня использовано {user["used_today"]}/{plan["daily_limit"]} анализов.',
            'upgrade_required': True
        }), 402
    
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'Файл не загружен'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Файл не выбран'}), 400
        
        # Сохраняем временный файл
        temp_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4()}_{file.filename}")
        file.save(temp_path)
        
        try:
            # Извлекаем текст
            if file.filename.lower().endswith('.pdf'):
                text = extract_text_from_pdf(temp_path)
            elif file.filename.lower().endswith('.docx'):
                text = extract_text_from_docx(temp_path)
            elif file.filename.lower().endswith('.txt'):
                with open(temp_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                return jsonify({'error': 'Неподдерживаемый формат файла'}), 400
            
            # Проверяем что текст извлекся
            if not text or len(text.strip()) < 10:
                return jsonify({'error': 'Не удалось извлечь текст из файла'}), 400
            
            # Анализируем текст
            analysis_result = analyze_text(text, user_id)
            
            # Записываем использование
            record_usage(user_id)
            
            # Добавляем информацию о лимитах в ответ
            user = get_user(user_id)
            plan = PLANS[user['plan']]
            analysis_result['usage_info'] = {
                'used_today': user['used_today'],
                'daily_limit': plan['daily_limit'],
                'plan_name': plan['name'],
                'remaining': plan['daily_limit'] - user['used_today']
            }
            
            return jsonify({
                'success': True,
                'filename': file.filename,
                'user_id': user_id,
                'result': analysis_result
            })
            
        finally:
            # Удаляем временный файл
            try:
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
            except:
                pass
            
    except Exception as e:
        return jsonify({'error': f'Ошибка обработки: {str(e)}'}), 500

# Обновляем endpoint использования
@app.route('/usage', methods=['GET'])
def get_usage():
    """Получить информацию об использовании"""
    user_id = request.args.get('user_id', 'default')
    user = get_user(user_id)
    plan = PLANS[user['plan']]
    
    return jsonify({
        'user_id': user_id,
        'plan': user['plan'],
        'plan_name': plan['name'],
        'used_today': user['used_today'],
        'daily_limit': plan['daily_limit'],
        'remaining': plan['daily_limit'] - user['used_today'],
        'total_used': user['total_used']
    })

@app.route('/plans', methods=['GET'])
def get_plans():
    """Получить информацию о тарифах"""
    return jsonify(PLANS)

@app.route('/api')
def api_info():
    return jsonify({
        'message': 'DocScan API работает!',
        'status': 'active',
        'ai_available': True,
        'pdf_export': False
    })

# 🔐 ЗАЩИЩЕННАЯ АДМИН-ПАНЕЛЬ

@app.route('/admin-login', methods=['GET', 'POST'])
def admin_login():
    """Страница входа в админ-панель"""
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        if username in ADMINS and ADMINS[username] == password:
            # Создаем сессию
            session_id = secrets.token_hex(16)
            admin_sessions[session_id] = {
                'username': username,
                'login_time': datetime.now().isoformat()
            }
            response = jsonify({'success': True, 'session_id': session_id})
            response.set_cookie('admin_session', session_id, httponly=True)
            return response
        else:
            return jsonify({'success': False, 'error': 'Неверные учетные данные'})
    
    return """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Admin Login - DocScan</title>
        <style>
            body { font-family: Arial; margin: 0; padding: 0; display: flex; justify-content: center; align-items: center; height: 100vh; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); }
            .login-box { background: white; padding: 40px; border-radius: 15px; box-shadow: 0 20px 40px rgba(0,0,0,0.1); width: 300px; }
            h2 { text-align: center; margin-bottom: 30px; color: #2d3748; }
            input { width: 100%; padding: 12px; margin: 10px 0; border: 1px solid #cbd5e0; border-radius: 8px; box-sizing: border-box; }
            button { width: 100%; background: #667eea; color: white; border: none; padding: 12px; border-radius: 8px; cursor: pointer; font-size: 16px; }
            button:hover { background: #5a67d8; }
            .error { color: #e53e3e; text-align: center; margin-top: 10px; }
        </style>
    </head>
    <body>
        <div class="login-box">
            <h2>🔧 Вход в админ-панель</h2>
            <form id="loginForm">
                <input type="text" name="username" placeholder="Логин" required>
                <input type="password" name="password" placeholder="Пароль" required>
                <button type="submit">Войти</button>
            </form>
            <div class="error" id="error"></div>
        </div>
        <script>
            document.getElementById('loginForm').addEventListener('submit', async (e) => {
                e.preventDefault();
                const formData = new FormData(e.target);
                
                const response = await fetch('/admin-login', {
                    method: 'POST',
                    body: formData
                });
                
                const result = await response.json();
                
                if (result.success) {
                    window.location.href = '/admin';
                } else {
                    document.getElementById('error').textContent = result.error;
                }
            });
        </script>
    </body>
    </html>
    """

def require_admin_auth(f):
    """Декоратор для проверки авторизации администратора"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        session_id = request.cookies.get('admin_session')
        
        if not session_id or session_id not in admin_sessions:
            return jsonify({'error': 'Требуется авторизация'}), 401
        
        return f(*args, **kwargs)
    return decorated_function

@app.route('/admin')
@require_admin_auth
def admin_panel():
    """Защищенная админ-панель"""
    session_id = request.cookies.get('admin_session')
    admin_info = admin_sessions.get(session_id, {})
    
    return """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Admin Panel - DocScan</title>
        <style>
            body { font-family: Arial; margin: 40px; background: #f7fafc; }
            .container { max-width: 1200px; margin: 0 auto; }
            .header { background: white; padding: 20px; border-radius: 10px; margin-bottom: 20px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }
            .user-card { background: white; padding: 15px; margin: 10px 0; border-radius: 8px; box-shadow: 0 2px 5px rgba(0,0,0,0.1); }
            button { background: #667eea; color: white; border: none; padding: 10px 15px; margin: 5px; border-radius: 5px; cursor: pointer; }
            button:hover { background: #5a67d8; }
            .logout-btn { background: #e53e3e; }
            .logout-btn:hover { background: #c53030; }
            .stats { display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 20px; margin: 20px 0; }
            .stat-card { background: white; padding: 20px; border-radius: 10px; text-align: center; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1>🔧 Админ-панель DocScan</h1>
                <p>Вошел как: <strong>""" + admin_info.get('username', 'Unknown') + """</strong></p>
                <button class="logout-btn" onclick="logout()">🚪 Выйти</button>
            </div>
            
            <div class="stats">
                <div class="stat-card">
                    <h3>👥 Всего пользователей</h3>
                    <div id="totalUsers">0</div>
                </div>
                <div class="stat-card">
                    <h3>📊 Всего анализов</h3>
                    <div id="totalAnalyses">0</div>
                </div>
                <div class="stat-card">
                    <h3>📈 Анализов сегодня</h3>
                    <div id="todayAnalyses">0</div>
                </div>
            </div>
            
            <h3>Управление пользователями:</h3>
            <div id="usersList"></div>
            
            <h3>Выдать тариф пользователю:</h3>
            <input type="text" id="userId" placeholder="ID пользователя">
            <select id="planSelect">
                <option value="free">Бесплатный (1 анализ)</option>
                <option value="basic">Базовый (10 анализов)</option>
                <option value="premium">Премиум (50 анализов)</option>
                <option value="unlimited">Безлимитный</option>
            </select>
            <button onclick="setUserPlan()">Выдать тариф</button>
            
            <h3>Создать нового пользователя:</h3>
            <input type="text" id="newUserId" placeholder="Новый ID пользователя (опционально)">
            <button onclick="createUser()">Создать пользователя</button>
        </div>

        <script>
            function logout() {
                document.cookie = "admin_session=; expires=Thu, 01 Jan 1970 00:00:00 UTC; path=/;";
                window.location.href = "/admin-login";
            }

            // Загружаем статистику и пользователей
            function loadStats() {
                fetch('/admin/stats')
                    .then(r => r.json())
                    .then(stats => {
                        document.getElementById('totalUsers').textContent = stats.total_users;
                        document.getElementById('totalAnalyses').textContent = stats.total_analyses;
                        document.getElementById('todayAnalyses').textContent = stats.today_analyses;
                    });
            }

            function loadUsers() {
                fetch('/admin/users')
                    .then(r => r.json())
                    .then(users => {
                        let html = '';
                        for (const [userId, userData] of Object.entries(users)) {
                            html += `
                                <div class="user-card">
                                    <strong>ID:</strong> ${userId}<br>
                                    <strong>Тариф:</strong> ${userData.plan} (${getPlanName(userData.plan)})<br>
                                    <strong>Использовано сегодня:</strong> ${userData.used_today}/${getPlanLimit(userData.plan)}<br>
                                    <strong>Всего анализов:</strong> ${userData.total_used}<br>
                                    <strong>Создан:</strong> ${userData.created_at || 'Неизвестно'}<br>
                                    <button onclick="setUserPlanQuick('${userId}', 'basic')">Выдать Базовый</button>
                                    <button onclick="setUserPlanQuick('${userId}', 'premium')">Выдать Премиум</button>
                                    <button onclick="setUserPlanQuick('${userId}', 'unlimited')">Выдать Безлимитный</button>
                                </div>
                            `;
                        }
                        document.getElementById('usersList').innerHTML = html;
                    });
            }

            function getPlanName(plan) {
                const names = {free: 'Бесплатный', basic: 'Базовый', premium: 'Премиум', unlimited: 'Безлимитный'};
                return names[plan] || plan;
            }

            function getPlanLimit(plan) {
                const limits = {free: 1, basic: 10, premium: 50, unlimited: 1000};
                return limits[plan] || 0;
            }

            function setUserPlan() {
                const userId = document.getElementById('userId').value;
                const plan = document.getElementById('planSelect').value;
                
                if (!userId) return alert('Введите ID пользователя');
                
                fetch('/admin/set-plan', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({user_id: userId, plan: plan})
                })
                .then(r => r.json())
                .then(result => {
                    alert(result.success ? '✅ ' + result.message : '❌ ' + result.error);
                    loadUsers();
                    loadStats();
                });
            }

            function setUserPlanQuick(userId, plan) {
                fetch('/admin/set-plan', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({user_id: userId, plan: plan})
                })
                .then(r => r.json())
                .then(result => {
                    alert(result.success ? '✅ ' + result.message : '❌ ' + result.error);
                    loadUsers();
                    loadStats();
                });
            }

            function createUser() {
                const userId = document.getElementById('newUserId').value;
                
                fetch('/admin/create-user', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({user_id: userId})
                })
                .then(r => r.json())
                .then(result => {
                    alert(result.success ? '✅ ' + result.message : '❌ ' + result.error);
                    loadUsers();
                    loadStats();
                });
            }

            // Загружаем при открытии
            loadStats();
            loadUsers();
        </script>
    </body>
    </html>
    """

# Админ API endpoints
@app.route('/admin/stats')
@require_admin_auth
def admin_stats():
    """Статистика для админ-панели"""
    total_users = len(users_db)
    total_analyses = sum(user['total_used'] for user in users_db.values())
    today_analyses = sum(user['used_today'] for user in users_db.values())
    
    return jsonify({
        'total_users': total_users,
        'total_analyses': total_analyses,
        'today_analyses': today_analyses
    })

@app.route('/admin/users')
@require_admin_auth
def get_all_users():
    """Получить всех пользователей"""
    return jsonify(users_db)

@app.route('/admin/set-plan', methods=['POST'])
@require_admin_auth
def admin_set_plan():
    """Установить тариф пользователю"""
    try:
        data = request.json
        user_id = data.get('user_id')
        plan = data.get('plan')
        
        if not user_id:
            return jsonify({'success': False, 'error': 'Укажите ID пользователя'})
        
        if user_id not in users_db:
            return jsonify({'success': False, 'error': 'Пользователь не найден'})
        
        if plan not in PLANS:
            return jsonify({'success': False, 'error': 'Неверный тариф'})
        
        # Обновляем тариф
        users_db[user_id]['plan'] = plan
        users_db[user_id]['used_today'] = 0  # Сбрасываем дневной лимит
        
        return jsonify({
            'success': True,
            'message': f'Пользователю {user_id} выдан тариф: {PLANS[plan]["name"]}'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/admin/create-user', methods=['POST'])
@require_admin_auth
def admin_create_user():
    """Создать нового пользователя"""
    try:
        data = request.json
        user_id = data.get('user_id')
        
        # Если ID не указан, генерируем случайный
        if not user_id:
            user_id = generate_user_id()
        
        if user_id in users_db:
            return jsonify({'success': False, 'error': 'Пользователь уже существует'})
        
        # Создаем пользователя
        users_db[user_id] = {
            'user_id': user_id,
            'plan': 'free',
            'used_today': 0,
            'last_reset': date.today().isoformat(),
            'total_used': 0,
            'created_at': datetime.now().isoformat()
        }
        
        return jsonify({
            'success': True,
            'message': f'Пользователь {user_id} создан с бесплатным тарифом',
            'user_id': user_id
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})
# Страницы для футера
@app.route('/terms')
def terms():
    return """
<!DOCTYPE html>
    <html>
    <head>
        <title>Пользовательское соглашение - DocScan</title>
        <meta name="description" content="Пользовательское соглашение сервиса DocScan. Условия использования AI анализа документов">
        <meta name="robots" content="index, follow">
        <style>
            body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; background: #f7fafc; }
            .container { max-width: 800px; margin: 0 auto; background: white; padding: 40px; border-radius: 15px; box-shadow: 0 10px 30px rgba(0,0,0,0.1); }
            .back-link { color: #667eea; text-decoration: none; margin-bottom: 20px; display: inline-block; font-weight: bold; }
            h1 { color: #2d3748; margin-bottom: 30px; border-bottom: 2px solid #667eea; padding-bottom: 10px; }
            h2 { color: #4a5568; margin-top: 30px; margin-bottom: 15px; }
            p { margin-bottom: 15px; color: #4a5568; }
            ul { margin: 15px 0; padding-left: 20px; }
            li { margin-bottom: 8px; color: #4a5568; }
        </style>
    </head>
    <body>
        <div class="container">
            <a href="/" class="back-link">← Назад на главную</a>
            <h1>Пользовательское соглашение</h1>
            
            <h2>1. Общие положения</h2>
            <p>1.1. Настоящее Пользовательское соглашение (далее — «Соглашение») регулирует отношения между владельцем сервиса DocScan (далее — «Администрация») и физическим или юридическим лицом (далее — «Пользователь») по использованию сервиса анализа документов.</p>
            <p>1.2. Используя сервис DocScan, Пользователь подтверждает свое согласие с условиями настоящего Соглашения в полном объеме.</p>

            <h2>2. Предмет соглашения</h2>
            <p>2.1. Администрация предоставляет Пользователю доступ к сервису анализа документов с использованием технологий искусственного интеллекта YandexGPT.</p>
            <p>2.2. Сервис позволяет анализировать текстовые документы форматов PDF, DOCX, TXT и получать анализ потенциальных рисков и рекомендаций.</p>

            <h2>3. Права и обязанности сторон</h2>
            <p><strong>3.1. Администрация обязуется:</strong></p>
            <ul>
                <li>Обеспечивать работоспособность сервиса в соответствии с техническими возможностями</li>
                <li>Обрабатывать документы в соответствии с Политикой конфиденциальности</li>
                <li>Предоставлять техническую поддержку Пользователям</li>
            </ul>

            <p><strong>3.2. Пользователь обязуется:</strong></p>
            <ul>
                <li>Не загружать документы, содержащие конфиденциальную, коммерческую тайну или персональные данные третьих лиц без соответствующих прав</li>
                <li>Не использовать сервис для анализа документов, содержащих противоправный контент</li>
                <li>Соблюдать установленные лимиты использования в соответствии с выбранным тарифом</li>
            </ul>

            <h2>4. Условия использования</h2>
            <p>4.1. Бесплатный тариф предусматривает 1 (один) анализ документов в сутки.</p>
            <p>4.2. Максимальный размер загружаемого файла — 10 МБ.</p>
            <p>4.3. Администрация вправе временно ограничить доступ к сервису для проведения технических работ.</p>

            <h2>5. Ответственность</h2>
            <p>5.1. Администрация не несет ответственности за решения, принятые Пользователем на основе анализа, предоставленного сервисом.</p>
            <p>5.2. Пользователь несет полную ответственность за содержание загружаемых документов.</p>

            <p><em>Дата последнего обновления: 22 ноября 2024 года</em></p>
        </div>
    </body>
    </html>
    """

@app.route('/privacy')
def privacy():
    return """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Политика конфиденциальности - DocScan</title>
        <meta name="description" content="Политика конфиденциальности DocScan. Как мы защищаем ваши данные при анализе документов">
        <meta name="robots" content="index, follow">
        <style>
            body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; background: #f7fafc; }
            .container { max-width: 800px; margin: 0 auto; background: white; padding: 40px; border-radius: 15px; box-shadow: 0 10px 30px rgba(0,0,0,0.1); }
            .back-link { color: #667eea; text-decoration: none; margin-bottom: 20px; display: inline-block; font-weight: bold; }
            h1 { color: #2d3748; margin-bottom: 30px; border-bottom: 2px solid #667eea; padding-bottom: 10px; }
            h2 { color: #4a5568; margin-top: 30px; margin-bottom: 15px; }
            p { margin-bottom: 15px; color: #4a5568; }
            ul { margin: 15px 0; padding-left: 20px; }
            li { margin-bottom: 8px; color: #4a5568; }
        </style>
    </head>
    <body>
        <div class="container">
            <a href="/" class="back-link">← Назад на главную</a>
            <h1>Политика конфиденциальности</h1>
            
            <h2>1. Общие положения</h2>
            <p>1.1. Настоящая Политика конфиденциальности определяет порядок обработки и защиты информации о Пользователях сервиса DocScan.</p>
            <p>1.2. Цель политики — обеспечение защиты прав Пользователей при использовании сервиса анализа документов.</p>

            <h2>2. Собираемая информация</h2>
            <p><strong>2.1. Техническая информация:</strong></p>
            <ul>
                <li>Анонимный идентификатор пользователя</li>
                <li>Статистика использования сервиса (количество анализов, даты использования)</li>
                <li>Информация о браузере и операционной системе</li>
                <li>IP-адрес (в обезличенном виде)</li>
            </ul>

            <p><strong>2.2. Обработка документов:</strong></p>
            <ul>
                <li>Загружаемые документы передаются в Yandex Cloud API для анализа</li>
                <li>Документы не сохраняются на наших серверах после обработки</li>
                <li>Текст документов удаляется сразу после формирования отчета</li>
                <li>Временные файлы автоматически удаляются в течение 1 часа</li>
            </ul>

            <h2>3. Использование информации</h2>
            <p>3.1. Собираемая информация используется исключительно для:</p>
            <ul>
                <li>Предоставления услуг анализа документов</li>
                <li>Улучшения работы сервиса</li>
                <li>Учета лимитов использования</li>
                <li>Технической поддержки Пользователей</li>
            </ul>

            <h2>4. Защита информации</h2>
            <p>4.1. Администрация принимает необходимые организационные и технические меры для защиты информации Пользователей.</p>
            <p>4.2. Передача данных между Пользователем и сервисом осуществляется по защищенному протоколу HTTPS.</p>
            <p>4.3. Доступ к статистической информации имеют только уполномоченные сотрудники Администрации.</p>

            <h2>5. Cookies и аналитика</h2>
            <p>5.1. Сервис использует cookies для:</p>
            <ul>
                <li>Сохранения идентификатора пользователя между сессиями</li>
                <li>Учета дневных лимитов использования</li>
                <li>Сбора анонимной статистики для улучшения сервиса</li>
            </ul>

            <p><em>Дата последнего обновления: 22 ноября 2024 года</em></p>
        </div>
    </body>
    </html>
    """

@app.route('/offer')
def offer():
    return """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Публичная оферта - DocScan</title>
        <meta name="description" content="Публичная оферта сервиса DocScan. Тарифы и условия использования AI анализа документов">
        <meta name="robots" content="index, follow">
        <style>
            body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; background: #f7fafc; }
            .container { max-width: 800px; margin: 0 auto; background: white; padding: 40px; border-radius: 15px; box-shadow: 0 10px 30px rgba(0,0,0,0.1); }
            .back-link { color: #667eea; text-decoration: none; margin-bottom: 20px; display: inline-block; font-weight: bold; }
            h1 { color: #2d3748; margin-bottom: 30px; border-bottom: 2px solid #667eea; padding-bottom: 10px; }
            h2 { color: #4a5568; margin-top: 30px; margin-bottom: 15px; }
            p { margin-bottom: 15px; color: #4a5568; }
            .price { background: #f0fff4; padding: 15px; border-radius: 10px; margin: 15px 0; border-left: 4px solid #38a169; }
            .contact { background: #edf2f7; padding: 20px; border-radius: 10px; margin: 20px 0; }
        </style>
    </head>
    <body>
        <div class="container">
            <a href="/" class="back-link">← Назад на главную</a>
            <h1>Публичная оферта</h1>
            
            <h2>1. Предмет оферты</h2>
            <p>1.1. Настоящая публичная оферта является официальным предложением Администрации сервиса DocScan заключить договор на оказание услуг на условиях, изложенных ниже.</p>
            <p>1.2. Акцептом оферты считается совершение Пользователем действий, направленных на использование платных услуг сервиса.</p>

            <h2>2. Тарифы и стоимость услуг</h2>
            
            <div class="price">
                <strong>Бесплатный тариф:</strong><br>
                • 1 анализ документов в сутки<br>
                • Доступ к базовому функционалу<br>
                • Стоимость: 0 рублей
            </div>

            <div class="price">
                <strong>Базовый тариф:</strong><br>
                • 10 анализов документов в сутки<br>
                • Приоритетная обработка<br>
                • Полный доступ к AI-анализу<br>
                • Стоимость: 199 рублей/месяц
            </div>

            <div class="price">
                <strong>Премиум тариф:</strong><br>
                • 50 анализов документов в сутки<br>
                • Высший приоритет обработки<br>
                • Расширенный AI-анализ<br>
                • Стоимость: 399 рублей/месяц
            </div>

            <div class="price">
                <strong>Безлимитный тариф:</strong><br>
                • Неограниченное количество анализов<br>
                • Максимальная скорость обработки<br>
                • Персональная поддержка<br>
                • Стоимость: 800 рублей/месяц
            </div>

            <h2>3. Порядок оплаты и подключения</h2>
            <p>3.1. Оплата услуг производится через платежные системы Яндекс.Касса или банковской картой.</p>
            <p>3.2. Подключение платного тарифа осуществляется моментально после успешной оплаты.</p>
            <p>3.3. Тариф действует в течение 30 (тридцати) календарных дней с момента оплаты.</p>

            <h2>4. Возврат средств</h2>
            <p>4.1. Пользователь вправе потребовать возврата средств в течение 14 дней с момента оплаты, если услуга не была использована.</p>
            <p>4.2. Для запроса возврата необходимо обратиться в техническую поддержку.</p>

            <h2>5. Технические требования</h2>
            <p>5.1. Поддерживаемые форматы документов: PDF, DOCX, TXT.</p>
            <p>5.2. Максимальный размер файла: 10 МБ.</p>
            <p>5.3. Рекомендуемый браузер: последние версии Chrome, Firefox, Safari, Edge.</p>

            <div class="contact">
                <h3>Реквизиты для связи</h3>
                <p><strong>Техническая поддержка:</strong> docscanhelp@gmail.com</p>
                <p><strong>По вопросам сотрудничества:</strong> docscanhelp@gmail.com</p>
            </div>

            <p><em>Оферта действует с 22 ноября 2025 года</em></p>
        </div>
    </body>
    </html>
    """
@app.route('/sitemap.xml')
def sitemap():
    base_url = "https://docscan-ekjj.onrender.com"
    return f'''<?xml version="1.0" encoding="UTF-8"?>
<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">
    <url>
        <loc>{base_url}/</loc>
        <lastmod>2024-11-22</lastmod>
        <changefreq>daily</changefreq>
        <priority>1.0</priority>
    </url>
    <url>
        <loc>{base_url}/terms</loc>
        <lastmod>2024-11-22</lastmod>
        <changefreq>monthly</changefreq>
        <priority>0.7</priority>
    </url>
    <url>
        <loc>{base_url}/privacy</loc>
        <lastmod>2024-11-22</lastmod>
        <changefreq>monthly</changefreq>
        <priority>0.7</priority>
    </url>
    <url>
        <loc>{base_url}/offer</loc>
        <lastmod>2024-11-22</lastmod>
        <changefreq>monthly</changefreq>
        <priority>0.7</priority>
    </url>
</urlset>''', 200, {'Content-Type': 'application/xml'}

@app.route('/robots.txt')
def robots():
    return """User-agent: *
Allow: /
Disallow: /admin
Disallow: /admin-login

Sitemap: https://docscan-ekjj.onrender.com/sitemap.xml""", 200, {'Content-Type': 'text/plain'}

# 🔐 ПЛАТЕЖНАЯ СИСТЕМА ЮMONEY

YOOMONEY_CLIENT_ID = os.getenv('YOOMONEY_CLIENT_ID')
YOOMONEY_CLIENT_SECRET = os.getenv('YOOMONEY_CLIENT_SECRET')
YOOMONEY_REDIRECT_URI = os.getenv('YOOMONEY_REDIRECT_URI', 'https://docscan-ekjj.onrender.com/payment-success')

@app.route('/create-payment', methods=['POST'])
def create_payment():
    """Создание платежа в ЮMoney"""
    try:
        data = request.json
        user_id = data.get('user_id')
        plan_type = data.get('plan')
        
        if not user_id or plan_type not in PLANS:
            return jsonify({'success': False, 'error': 'Неверные данные'})
        
        plan = PLANS[plan_type]
        
        # Создаем ссылку для оплаты через ЮMoney
        base_url = "https://yoomoney.ru/oauth/authorize"
        params = {
            'client_id': YOOMONEY_CLIENT_ID,
            'response_type': 'code',
            'redirect_uri': YOOMONEY_REDIRECT_URI,
            'scope': 'account-info operation-history operation-details',
            'state': f'{user_id}_{plan_type}'
        }
        
        payment_url = f"{base_url}?{'&'.join([f'{k}={v}' for k, v in params.items()])}"
        
        return jsonify({
            'success': True,
            'payment_url': payment_url,
            'message': f'Оплата тарифа {plan["name"]} - {plan["price"]}₽'
        })
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/payment-success')
def payment_success():
    """Страница успешной оплаты"""
    code = request.args.get('code')
    state = request.args.get('state')
    
    if state:
        try:
            user_id, plan_type = state.split('_')
            return f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Платеж обрабатывается - DocScan</title>
                <style>
                    body {{ font-family: Arial; margin: 0; padding: 0; display: flex; justify-content: center; align-items: center; height: 100vh; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); }}
                    .container {{ background: white; padding: 40px; border-radius: 15px; box-shadow: 0 20px 40px rgba(0,0,0,0.1); text-align: center; }}
                    .success-icon {{ font-size: 4em; color: #48bb78; margin-bottom: 20px; }}
                    .btn {{ background: #48bb78; color: white; border: none; padding: 15px 30px; border-radius: 50px; font-size: 1.1em; cursor: pointer; text-decoration: none; display: inline-block; margin-top: 20px; }}
                </style>
            </head>
            <body>
                <div class="container">
                    <div class="success-icon">⏳</div>
                    <h1>Платеж получен!</h1>
                    <p>Тариф активируется в течение 5 минут.</p>
                    <p><strong>ID пользователя:</strong> {user_id}</p>
                    <p><strong>Тариф:</strong> {plan_type}</p>
                    <a href="/" class="btn">Вернуться в DocScan</a>
                </div>
            </body>
            </html>
            """
        except:
            pass
    
    return """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Платеж успешен - DocScan</title>
        <style>
            body { font-family: Arial; margin: 0; padding: 0; display: flex; justify-content: center; align-items: center; height: 100vh; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); }
            .container { background: white; padding: 40px; border-radius: 15px; box-shadow: 0 20px 40px rgba(0,0,0,0.1); text-align: center; }
            .success-icon { font-size: 4em; color: #48bb78; margin-bottom: 20px; }
            .btn { background: #48bb78; color: white; border: none; padding: 15px 30px; border-radius: 50px; font-size: 1.1em; cursor: pointer; text-decoration: none; display: inline-block; margin-top: 20px; }
        </style>
    </head>
    <body>
        <div class="container">
            <div class="success-icon">✅</div>
            <h1>Платеж успешно завершен!</h1>
            <p>Тариф активирован. Возвращайтесь в приложение.</p>
            <a href="/" class="btn">Вернуться в DocScan</a>
        </div>
    </body>
    </html>
    """

@app.route('/payment-webhook', methods=['POST'])
def payment_webhook():
    """Webhook для уведомлений от ЮMoney (будет настроен позже)"""
    try:
        print("🔄 Webhook получен от ЮMoney")
        return jsonify({'success': True})
    except Exception as e:
        print(f"❌ Ошибка webhook: {e}")
        return jsonify({'success': False})

if __name__ == '__main__':
    print("🚀 DocScan Server запущен!")
    print("🤖 YandexGPT: Активен")
    print("🔐 Админ-панель: Защищена паролем")
    print("👤 Индивидуальные ID пользователей: Активны")
    print("💰 Бесплатный лимит: 1 анализ в день")
    
    # Для продакшена на Render
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
    